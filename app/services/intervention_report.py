"""Builds the downloadable School-Based Intervention Plan .docx from an
InterventionRecommendation (Agent 2's output). Built programmatically with
python-docx (no external template dependency) so a fresh project doesn't
need a matching .docx template shipped alongside it. Fixed headings/labels
come from the DB-backed translations table (app/services/i18n_lookup.py)
instead of a hardcoded LABELS dict, keyed by the same language code used
elsewhere in the app."""
import io

from docx import Document
from docx.shared import Pt, RGBColor
from sqlalchemy.orm import Session

from app.schemas.risk import InterventionRecommendation
from app.services.i18n_lookup import get_translation

STRATEGY_AREAS = ["Classroom", "Instruction", "Differentiated Learning", "Behaviour", "Movement", "Counselling", "Parents"]

_HEADER_BLUE = RGBColor(0x1A, 0x3A, 0xAD)


def _label(key: str, lang_code: str, db: Session, default: str) -> str:
    return get_translation(key, lang_code, db, default=default)


def _add_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = _HEADER_BLUE


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def generate_intervention_docx(plan: InterventionRecommendation, lang_code: str, db: Session) -> bytes:
    doc = Document()

    title = _label("report.school_based_intervention_plan", lang_code, db, "School-Based Intervention Plan")
    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = _HEADER_BLUE

    disclaimer_p = doc.add_paragraph(plan.disclaimer)
    disclaimer_p.runs[0].italic = True
    disclaimer_p.runs[0].font.size = Pt(9)

    _add_heading(doc, "Student Information")
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Light Grid Accent 1"
    for label, value in [
        ("Student Name", plan.student_name),
        ("Class", plan.class_name),
        ("Age", str(plan.age) if plan.age is not None else "-"),
        ("Risk Level", plan.risk_level),
        ("School", plan.school_name),
        ("Case Reference", plan.case_reference),
        ("Prepared By", plan.prepared_by),
        ("Date", plan.date),
    ]:
        row = info_table.add_row().cells
        row[0].text = label
        row[1].text = value

    _add_heading(doc, "Reason for Intervention")
    _add_bullets(doc, plan.reason_for_intervention)

    _add_heading(doc, "Intervention Objectives")
    _add_bullets(doc, plan.intervention_objectives)

    _add_heading(doc, "AI Recommended Intervention Plan")
    strategy_table = doc.add_table(rows=1, cols=5)
    strategy_table.style = "Light Grid Accent 1"
    header_cells = strategy_table.rows[0].cells
    for i, col in enumerate(["Area", "Strategy", "Responsible Person", "Frequency", "Success Indicator"]):
        header_cells[i].text = col
    by_area = {s.area: s for s in plan.strategies}
    for area in STRATEGY_AREAS:
        s = by_area.get(area)
        row = strategy_table.add_row().cells
        row[0].text = area
        row[1].text = s.strategy if s else "-"
        row[2].text = s.responsible if s else "-"
        row[3].text = s.frequency if s else "-"
        row[4].text = s.success_indicator if s else "-"

    _add_heading(doc, "Recommended Tools")
    _add_bullets(doc, plan.recommended_tools)

    _add_heading(doc, "Parent Support Guide")
    _add_bullets(doc, plan.home_strategies)

    _add_heading(doc, "Expected Outcomes")
    _add_bullets(doc, plan.expected_outcomes)

    _add_heading(doc, "Monitoring Checklist")
    _add_bullets(doc, plan.monitoring_checklist)

    _add_heading(doc, "Counselor Recommendation")
    doc.add_paragraph(plan.counselor_recommendation)

    if plan.referral_recommended:
        _add_heading(doc, "Referral Recommendation")
        doc.add_paragraph(plan.referral_reason or "Referral to Agent 3 recommended.")

    _add_heading(doc, "Action Items")
    doc.add_paragraph("Teacher:", style="Intense Quote")
    _add_bullets(doc, plan.action_items_teacher)
    doc.add_paragraph("Counselor:", style="Intense Quote")
    _add_bullets(doc, plan.action_items_counselor)
    doc.add_paragraph("Parent:", style="Intense Quote")
    _add_bullets(doc, plan.action_items_parent)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
